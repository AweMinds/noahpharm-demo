from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
import os
import logging
from llm_service import get_llm_service

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)  # 允许跨域请求

# 配置允许访问的基础路径（出于安全考虑）
ALLOWED_BASE_PATHS = [
    'E:\\temp\\氨氯地平缬沙坦氢氯噻嗪片-demo'
]

def is_path_allowed(path):
    """检查路径是否在允许的范围内"""
    normalized_path = os.path.normpath(path)
    for allowed_path in ALLOWED_BASE_PATHS:
        if normalized_path.startswith(os.path.normpath(allowed_path)):
            return True
    return False

@app.route('/api/folders', methods=['GET'])
def get_subfolders():
    """获取指定路径下的子文件夹"""
    try:
        folder_path = request.args.get('path')
        if not folder_path:
            return jsonify({'error': '缺少path参数'}), 400
        
        # 安全检查
        if not is_path_allowed(folder_path):
            return jsonify({'error': '不允许访问此路径'}), 403
        
        # 检查路径是否存在
        if not os.path.exists(folder_path):
            logger.warning(f"路径不存在: {folder_path}")
            return jsonify({'subfolders': []})
        
        # 检查是否为目录
        if not os.path.isdir(folder_path):
            return jsonify({'error': '指定路径不是目录'}), 400
        
        # 读取子文件夹
        subfolders = []
        try:
            for item in os.listdir(folder_path):
                item_path = os.path.join(folder_path, item)
                if os.path.isdir(item_path):
                    subfolders.append(item)
        except PermissionError:
            return jsonify({'error': '没有权限访问此目录'}), 403
        
        logger.info(f"成功读取路径 {folder_path}, 找到 {len(subfolders)} 个子文件夹")
        return jsonify({'subfolders': sorted(subfolders)})
        
    except Exception as e:
        logger.error(f"读取文件夹时发生错误: {str(e)}")
        return jsonify({'error': f'服务器内部错误: {str(e)}'}), 500

@app.route('/api/check-path', methods=['GET'])
def check_path_exists():
    """检查路径是否存在"""
    try:
        folder_path = request.args.get('path')
        if not folder_path:
            return jsonify({'error': '缺少path参数'}), 400
        
        # 安全检查
        if not is_path_allowed(folder_path):
            return jsonify({'exists': False, 'reason': '路径不在允许范围内'})
        
        exists = os.path.exists(folder_path) and os.path.isdir(folder_path)
        return jsonify({'exists': exists})
        
    except Exception as e:
        logger.error(f"检查路径时发生错误: {str(e)}")
        return jsonify({'error': f'服务器内部错误: {str(e)}'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    """健康检查接口"""
    return jsonify({'status': 'healthy', 'message': '文件系统API服务正常运行'})

@app.route('/api/folders/all', methods=['GET'])
def get_all_folders():
    """获取所有配置路径的子文件夹"""
    try:
        # 定义所有需要读取的路径
        paths = {
            'CDE同类品种-临床备案公示平台试验信息': 'E:\\temp\\氨氯地平缬沙坦氢氯噻嗪片-demo\\CDE同类品种-临床备案公示平台试验信息',
            '国外试验文献调研': 'E:\\temp\\氨氯地平缬沙坦氢氯噻嗪片-demo\\国外试验文献调研',
            '法规_指导原则_用药指南': 'E:\\temp\\氨氯地平缬沙坦氢氯噻嗪片-demo\\法规_指导原则_用药指南',
            '说明书': 'E:\\temp\\氨氯地平缬沙坦氢氯噻嗪片-demo\\说明书'
        }
        
        result = {}
        for section_name, folder_path in paths.items():
            if not is_path_allowed(folder_path):
                result[section_name] = []
                continue
            
            if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
                result[section_name] = []
                continue
            
            try:
                subfolders = []
                for item in os.listdir(folder_path):
                    item_path = os.path.join(folder_path, item)
                    if os.path.isdir(item_path):
                        subfolders.append(item)
                result[section_name] = sorted(subfolders)
            except PermissionError:
                logger.warning(f"没有权限访问目录: {folder_path}")
                result[section_name] = []
        
        logger.info(f"成功读取所有文件夹，共返回 {len(result)} 个区域的数据")
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"读取所有文件夹时发生错误: {str(e)}")
        return jsonify({'error': f'服务器内部错误: {str(e)}'}), 500

@app.route('/api/extract-info', methods=['POST'])
def extract_key_info():
    """提取关键信息API"""
    try:
        data = request.get_json()
        if not data or 'selected_items' not in data:
            return jsonify({'error': '缺少selected_items参数'}), 400
        
        selected_items = data['selected_items']
        if not selected_items:
            return jsonify({'error': '没有选择任何项目'}), 400
        
        # 定义基础路径映射
        base_paths = {
            'CDE同类品种-临床备案公示平台试验信息': 'E:\\temp\\氨氯地平缬沙坦氢氯噻嗪片-demo\\CDE同类品种-临床备案公示平台试验信息',
            '国外试验文献调研': 'E:\\temp\\氨氯地平缬沙坦氢氯噻嗪片-demo\\国外试验文献调研'
        }
        
        results = []
        
        # 获取LLM服务实例
        try:
            llm_service = get_llm_service()
        except Exception as e:
            return jsonify({'error': f'LLM服务初始化失败: {str(e)}'}), 500
        
        for item in selected_items:
            try:
                # 解析section和literature_name
                section_name, literature_name = item.split('/', 1)
                
                # 检查是否为支持的类型
                if section_name not in ['CDE同类品种-临床备案公示平台试验信息', '国外试验文献调研']:
                    continue
                
                # 构建文献文件夹路径
                base_path = base_paths[section_name]
                literature_folder = os.path.join(base_path, literature_name)
                
                # 查找MD文件
                md_file_path = llm_service.find_md_file_path(literature_folder, literature_name)
                
                # 读取MD文件内容
                content = llm_service.read_md_file(md_file_path)
                
                # 使用LLM提取关键信息，传入文献类型
                extracted_info = llm_service.extract_key_info(content, section_name)
                
                # 添加文献标识信息
                extracted_info['literature_name'] = literature_name
                extracted_info['section_name'] = section_name
                extracted_info['md_file_path'] = md_file_path
                
                results.append(extracted_info)
                
                logger.info(f"成功提取文献 {literature_name} ({section_name}) 的关键信息")
                
            except Exception as e:
                error_msg = f"处理文献 {item} 时发生错误: {str(e)}"
                logger.error(error_msg)
                results.append({
                    'literature_name': item,
                    'error': error_msg
                })
        
        return jsonify({
            'success': True,
            'results': results,
            'total_processed': len(results)
        })
        
    except Exception as e:
        logger.error(f"提取关键信息时发生错误: {str(e)}")
        return jsonify({'error': f'服务器内部错误: {str(e)}'}), 500

@app.route('/api/generate-summary', methods=['POST'])
def generate_summary():
    """生成方案摘要API"""
    try:
        data = request.get_json()
        if not data or 'literature_info' not in data:
            return jsonify({'error': '缺少literature_info参数'}), 400
        
        literature_info = data['literature_info']
        if not literature_info.get('md_file_path'):
            return jsonify({'error': '缺少md_file_path信息'}), 400
        
        # 获取LLM服务实例
        try:
            llm_service = get_llm_service()
        except Exception as e:
            return jsonify({'error': f'LLM服务初始化失败: {str(e)}'}), 500
        
        # 读取MD文件内容
        md_file_path = literature_info['md_file_path']
        content = llm_service.read_md_file(md_file_path)
        
        # 使用LLM生成方案摘要
        summary = llm_service.generate_summary(content)
        
        # 添加文献信息
        summary['literature_name'] = literature_info.get('literature_name', '')
        summary['section_name'] = literature_info.get('section_name', '')
        
        logger.info(f"成功生成文献 {literature_info.get('literature_name', '')} 的方案摘要")
        
        return jsonify({
            'success': True,
            'summary': summary
        })
        
    except Exception as e:
        logger.error(f"生成方案摘要时发生错误: {str(e)}")
        return jsonify({'error': f'服务器内部错误: {str(e)}'}), 500

@app.route('/api/download-summary', methods=['POST'])
def download_summary():
    """下载方案摘要Word文档API"""
    try:
        from io import BytesIO
        import tempfile
        
        data = request.get_json()
        if not data or 'summary' not in data:
            return jsonify({'error': '缺少summary参数'}), 400
        
        summary = data['summary']
        
        # 生成Word文档
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            return jsonify({'error': '缺少python-docx依赖，请安装: pip install python-docx'}), 500
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('关键文献方案摘要', 0)
        
        # 添加基本信息表格
        basic_info = [
            ('研究题目', summary.get('study_title', '')),
            ('申办者', summary.get('sponsor', '')),
            ('方案编号', summary.get('protocol_number', '')),
            ('组长单位/主要研究者', summary.get('principal_investigator', '')),
            ('试验药物名称及规格', summary.get('drug_name_specification', '')),
            ('适应症', summary.get('indication', '')),
            ('试验分期', summary.get('trial_phase', '')),
            ('研究中心数', summary.get('center_count', '')),
            ('研究周期', summary.get('study_period', ''))
        ]
        
        table = doc.add_table(rows=len(basic_info), cols=2)
        table.style = 'Table Grid'
        
        for i, (label, value) in enumerate(basic_info):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        # 添加研究目的
        doc.add_heading('研究目的', level=1)
        objectives = summary.get('study_objectives', {})
        doc.add_paragraph(f"主要目的: {objectives.get('primary_objective', '')}")
        doc.add_paragraph(f"次要目的: {objectives.get('secondary_objective', '')}")
        
        # 添加疗效指标
        doc.add_heading('疗效指标', level=1)
        endpoints = summary.get('efficacy_endpoints', {})
        doc.add_paragraph(f"主要终点: {endpoints.get('primary_endpoint', '')}")
        doc.add_paragraph(f"次要终点: {endpoints.get('secondary_endpoint', '')}")
        doc.add_paragraph(f"探索性终点: {endpoints.get('exploratory_endpoint', '')}")
        doc.add_paragraph(f"安全性评价: {endpoints.get('safety_evaluation', '')}")
        
        # 添加试验设计
        doc.add_heading('试验设计', level=1)
        design = summary.get('trial_design', {})
        doc.add_paragraph(f"研究人群选择及导入期设计依据: {design.get('study_population_selection', '')}")
        doc.add_paragraph(f"阳性对照药品选择及依据: {design.get('positive_control_selection', '')}")
        doc.add_paragraph(f"主要疗效终点的选择及依据: {design.get('primary_endpoint_selection', '')}")
        
        # 添加其他信息
        other_info = [
            ('试验流程', summary.get('trial_process', '')),
            ('样本量', summary.get('sample_size', '')),
            ('试验用药品，规格，用法用量', summary.get('investigational_drug', '')),
            ('合并治疗', summary.get('concomitant_treatment', '')),
            ('挽救治疗', summary.get('rescue_treatment', '')),
            ('入排标准', summary.get('inclusion_exclusion_criteria', '')),
            ('退出和中止/终止标准', summary.get('withdrawal_termination_criteria', '')),
            ('统计分析', summary.get('statistical_analysis', ''))
        ]
        
        for label, value in other_info:
            doc.add_heading(label, level=1)
            doc.add_paragraph(value)
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        # 返回文件
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name='方案摘要.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"下载方案摘要时发生错误: {str(e)}")
        return jsonify({'error': f'服务器内部错误: {str(e)}'}), 500

if __name__ == '__main__':
    print("启动文件系统API服务...")
    print("访问地址: http://localhost:5000")
    print("健康检查: http://localhost:5000/api/health")
    app.run(debug=True, host='0.0.0.0', port=5000) 